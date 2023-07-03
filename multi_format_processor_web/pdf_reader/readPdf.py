import sys
import os

import pdfminer
import pandas as pd
from docx import Document

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams



def read_pdf(pdf_path):

    fp = open(pdf_path, 'rb')

    parser = PDFParser(fp)
    doc = PDFDocument(parser)
    parser.set_document(doc)
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    
    text = ''
    
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)
        layout = device.get_result()
        
        for x in layout:
            if (isinstance(x, LTTextBoxHorizontal)):
                results = x.get_text()
                text += results
                
    return text

def save_as_txt(text):
    with open("output.txt", "w") as file:
        file.write(text)
    return

def save_as_word(text):
    document = Document()
    document.add_paragraph(text)
    document.save("output.docx")
    return